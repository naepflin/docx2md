#!/usr/bin/env python3

import sys
import os
import re
import base64
import mimetypes

import mammoth
import mistune
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def convert_image(image):
    """Convert embedded images to inline base64 markdown images."""
    with image.open() as img_bytes:
        data = img_bytes.read()
    content_type = image.content_type or "image/png"
    ext = content_type.split("/")[-1]
    b64 = base64.b64encode(data).decode("ascii")
    alt = image.alt_text or "image"
    src = f"data:{content_type};base64,{b64}"
    return {"src": src, "alt": alt, "ext": ext}


def save_image(image, output_dir, image_count):
    """Save embedded image to a file and return markdown reference."""
    with image.open() as img_bytes:
        data = img_bytes.read()
    content_type = image.content_type or "image/png"
    ext = content_type.split("/")[-1]
    if ext == "jpeg":
        ext = "jpg"
    alt = image.alt_text or "image"
    filename = f"image{image_count[0]}.{ext}"
    image_count[0] += 1
    filepath = os.path.join(output_dir, filename)
    with open(filepath, "wb") as f:
        f.write(data)
    return {"src": filename, "alt": alt}


def _apply_inline_formatting(md):
    """Apply the inline-level HTML → markdown substitutions."""
    md = re.sub(r"<strong>(.*?)</strong>", r"**\1**", md, flags=re.DOTALL)
    md = re.sub(r"<em>(.*?)</em>", r"*\1*", md, flags=re.DOTALL)
    md = re.sub(r'<a href="(.*?)">(.*?)</a>', r"[\2](\1)", md, flags=re.DOTALL)
    md = re.sub(r'<img src="(.*?)" alt="(.*?)" />', r"![\2](\1)", md)
    md = re.sub(r'<img src="(.*?)" />', r"![image](\1)", md)
    md = re.sub(r"<code>(.*?)</code>", r"`\1`", md, flags=re.DOTALL)
    md = re.sub(r"<sup>(.*?)</sup>", r"^\1^", md, flags=re.DOTALL)
    md = re.sub(r"<sub>(.*?)</sub>", r"~\1~", md, flags=re.DOTALL)
    return md


def _unescape_entities(text):
    """Turn the HTML entities mammoth emits back into literal characters."""
    text = text.replace("&lt;", "<")
    text = text.replace("&gt;", ">")
    text = text.replace("&quot;", '"')
    text = text.replace("&#39;", "'")
    text = text.replace("&nbsp;", " ")
    # &amp; last, so "&amp;lt;" does not become "<"
    text = text.replace("&amp;", "&")
    return text


def html_to_markdown(html):
    """Convert mammoth's HTML output to clean markdown."""
    # Tables are converted first and parked behind placeholders: their cells
    # hold block markup (paragraphs, lists, breaks) that the block-level rules
    # below would turn into newlines, and a newline inside a markdown table
    # cell ends the row.
    tables = []

    def stash_table(match):
        tables.append(table_to_markdown(match.group(0)))
        return f"\n\n\x00TABLE{len(tables) - 1}\x00\n\n"

    md = re.sub(r"<table\b[^>]*>.*?</table>", stash_table, html, flags=re.DOTALL)

    # Headings
    for level in range(6, 0, -1):
        prefix = "#" * level
        md = re.sub(
            rf"<h{level}>(.*?)</h{level}>",
            rf"{prefix} \1\n",
            md,
            flags=re.DOTALL,
        )

    # Line breaks
    md = re.sub(r"<br />", "\n", md)

    # Lists - unordered
    md = re.sub(r"<ul>", "\n", md)
    md = re.sub(r"</ul>", "\n", md)
    md = re.sub(r"<li>(.*?)</li>", r"- \1\n", md, flags=re.DOTALL)

    # Lists - ordered
    md = re.sub(r"<ol>", "\n", md)
    md = re.sub(r"</ol>", "\n", md)

    # Handle ordered list items with numbering
    counter = [0]

    def replace_ol_item(match):
        counter[0] += 1
        return f"{counter[0]}. {match.group(1)}"

    md = re.sub(r"<li>(.*?)</li>", replace_ol_item, md, flags=re.DOTALL)

    # Paragraphs
    md = re.sub(r"<p>(.*?)</p>", r"\1\n\n", md, flags=re.DOTALL)

    # Blockquotes
    md = re.sub(
        r"<blockquote>(.*?)</blockquote>",
        lambda m: "\n".join("> " + line for line in m.group(1).strip().split("\n")),
        md,
        flags=re.DOTALL,
    )

    # Code
    md = re.sub(
        r"<pre>(.*?)</pre>", r"```\n\1\n```\n", md, flags=re.DOTALL
    )

    # Bold, italic, links, images, inline code, super/subscript
    md = _apply_inline_formatting(md)

    # Strip any remaining HTML tags
    md = re.sub(r"<[^>]+>", "", md)

    md = _unescape_entities(md)

    # Clean up whitespace: collapse 3+ newlines to 2
    md = re.sub(r"\n{3,}", "\n\n", md)
    md = md.strip() + "\n"

    # Put the tables back now that no block-level rule can reach into them
    md = re.sub(r"\x00TABLE(\d+)\x00", lambda m: tables[int(m.group(1))], md)

    return md


def cell_to_markdown(cell_html):
    """Convert the inner HTML of one table cell to single-line markdown."""
    # Every block boundary inside the cell becomes a hard break, because a
    # markdown cell cannot span lines.
    sep = "\x00BR\x00"
    md = re.sub(r"<li[^>]*>(.*?)</li>", rf"{sep}• \1", cell_html, flags=re.DOTALL)
    md = re.sub(r"<br\s*/?>", sep, md)
    md = re.sub(r"</?(?:p|div|ul|ol|h[1-6]|blockquote)\b[^>]*>", sep, md)

    md = _apply_inline_formatting(md)
    md = re.sub(r"<[^>]+>", "", md)
    md = _unescape_entities(md)

    # Drop the empty blocks that Word's spacer paragraphs leave behind
    blocks = [re.sub(r"\s+", " ", b).strip() for b in md.split(sep)]
    text = "<br>".join(b for b in blocks if b)

    # A literal pipe would end the cell
    return text.replace("|", "\\|")


def table_to_markdown(table_html):
    """Convert one HTML table to a markdown table."""
    rows = re.findall(r"<tr\b[^>]*>(.*?)</tr>", table_html, re.DOTALL)
    if not rows:
        return ""

    # Each row is a list of cells; a colspan cell is followed by the empty
    # cells it swallowed, since markdown has no way to merge columns.
    parsed = []
    for row in rows:
        cells = []
        for _tag, attrs, inner in re.findall(
            r"<t([dh])\b([^>]*)>(.*?)</t\1>", row, re.DOTALL
        ):
            cells.append(cell_to_markdown(inner))
            span = re.search(r'colspan="(\d+)"', attrs)
            if span:
                cells.extend([""] * (int(span.group(1)) - 1))
        parsed.append(cells)

    num_cols = max(len(cells) for cells in parsed)
    if num_cols == 0:
        return ""

    md_rows = []
    for cells in parsed:
        cells = cells + [""] * (num_cols - len(cells))
        md_rows.append("| " + " | ".join(cells) + " |")

    separator = "| " + " | ".join(["---"] * num_cols) + " |"
    md_rows.insert(1, separator)

    return "\n".join(md_rows)


def convert(input_path, extract_images=False, image_dir=None):
    """Convert a .docx file to markdown string."""
    image_count = [1]

    if extract_images:
        img_dir = image_dir or os.path.splitext(input_path)[0] + "_images"
        os.makedirs(img_dir, exist_ok=True)

        def handle_image(image):
            info = save_image(image, img_dir, image_count)
            return {"src": info["src"]}

        convert_image_fn = mammoth.images.img_element(handle_image)
    else:
        def handle_image_inline(image):
            info = convert_image(image)
            return {"src": info["src"]}

        convert_image_fn = mammoth.images.img_element(handle_image_inline)

    with open(input_path, "rb") as f:
        result = mammoth.convert_to_html(f, convert_image=convert_image_fn)

    if result.messages:
        for msg in result.messages:
            print(f"Warning: {msg.message}", file=sys.stderr)

    return html_to_markdown(result.value)


HEADING_MAP = {
    1: "Heading 1",
    2: "Heading 2",
    3: "Heading 3",
    4: "Heading 4",
    5: "Heading 5",
    6: "Heading 6",
}


def _extract_text(tokens):
    """Recursively extract plain text from a list of inline tokens."""
    parts = []
    for tok in tokens:
        if "raw" in tok:
            parts.append(tok["raw"])
        elif tok.get("children"):
            parts.append(_extract_text(tok["children"]))
    return "".join(parts)


def _add_inline(paragraph, tokens):
    """Render inline markdown tokens into a docx paragraph."""
    for tok in tokens:
        ttype = tok["type"]
        if ttype == "text":
            paragraph.add_run(tok["raw"])
        elif ttype == "codespan":
            run = paragraph.add_run(tok["raw"])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif ttype == "strong":
            run = paragraph.add_run(_extract_text(tok.get("children", [])))
            run.bold = True
        elif ttype == "emphasis":
            run = paragraph.add_run(_extract_text(tok.get("children", [])))
            run.italic = True
        elif ttype == "link":
            text = _extract_text(tok["children"]) if tok.get("children") else tok.get("link", "")
            paragraph.add_run(text)
        elif ttype == "image":
            paragraph.add_run(f'[image: {tok.get("alt", "")}]')
        elif ttype == "softbreak":
            # A single newline inside a paragraph is a soft break: per the
            # Markdown spec it separates source lines only and renders as a
            # space, not a line break.
            paragraph.add_run(" ")
        elif ttype == "linebreak":
            # A hard break (trailing two spaces or a backslash) becomes an
            # actual line break within the same Word paragraph.
            paragraph.add_run().add_break()
        elif ttype == "inline_html":
            # Table cells carry <br> for the block breaks a markdown cell
            # cannot express; any other stray tag is markup, not content.
            raw = tok.get("raw", "")
            if re.match(r"<br\s*/?>", raw, re.IGNORECASE):
                paragraph.add_run().add_break()
            elif not raw.startswith("<"):
                paragraph.add_run(raw)
        else:
            if "children" in tok and tok["children"]:
                _add_inline(paragraph, tok["children"])
            elif "raw" in tok:
                paragraph.add_run(tok["raw"])


def _add_table(doc, token):
    """Render a markdown table token into a docx table."""
    children = token.get("children", [])
    head = next((c for c in children if c["type"] == "table_head"), None)
    body = next((c for c in children if c["type"] == "table_body"), None)

    head_cells = head.get("children", []) if head else []
    body_rows = body.get("children", []) if body else []

    num_cols = len(head_cells)
    if num_cols == 0:
        return
    num_rows = 1 + len(body_rows)
    table = doc.add_table(rows=num_rows, cols=num_cols, style="Table Grid")

    for ci, cell_tok in enumerate(head_cells):
        cell = table.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        _add_inline(p, cell_tok.get("children", []))
        for run in p.runs:
            run.bold = True

    for ri, row_tok in enumerate(body_rows):
        row_cells = row_tok.get("children", [])
        for ci, cell_tok in enumerate(row_cells):
            if ci >= num_cols:
                break
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_inline(p, cell_tok.get("children", []))


def _add_list(doc, tok):
    """Recursively render a list token into docx paragraphs."""
    ordered = tok["attrs"].get("ordered", False)
    depth = tok["attrs"].get("depth", 0)
    level = depth + 1
    suffix = f" {level}" if level > 1 else ""
    style = f"List Number{suffix}" if ordered else f"List Bullet{suffix}"

    for item in tok.get("children", []):
        for child in item.get("children", []):
            ctype = child["type"]
            if ctype in ("paragraph", "block_text"):
                p = doc.add_paragraph(style=style)
                _add_inline(p, child.get("children", []))
            elif ctype == "list":
                _add_list(doc, child)
            elif ctype == "block_code":
                p = doc.add_paragraph()
                run = p.add_run(child.get("raw", "").rstrip("\n"))
                run.font.name = "Courier New"
                run.font.size = Pt(9)


def convert_md_to_docx(md_text, input_path=None):
    """Convert markdown text to a python-docx Document."""
    md = mistune.create_markdown(renderer=None, plugins=["table", "strikethrough"])
    tokens = md(md_text)

    doc = Document()
    input_dir = os.path.dirname(os.path.abspath(input_path)) if input_path else None

    for tok in tokens:
        ttype = tok["type"]

        if ttype == "heading":
            level = tok["attrs"]["level"]
            p = doc.add_paragraph(style=HEADING_MAP.get(level, "Heading 1"))
            _add_inline(p, tok["children"])

        elif ttype == "paragraph":
            p = doc.add_paragraph()
            children = tok.get("children", [])
            if len(children) == 1 and children[0]["type"] == "image":
                img = children[0]
                src = img.get("src", "")
                if input_dir and not src.startswith(("http://", "https://", "data:")):
                    img_path = os.path.join(input_dir, src)
                    if os.path.exists(img_path):
                        doc.add_picture(img_path, width=Inches(5))
                        continue
                if src.startswith("data:"):
                    match = re.match(r"data:[^;]+;base64,(.+)", src)
                    if match:
                        import tempfile
                        data = base64.b64decode(match.group(1))
                        content_type = src.split(";")[0].split(":")[1]
                        ext = mimetypes.guess_extension(content_type) or ".png"
                        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                            tmp.write(data)
                            tmp_path = tmp.name
                        try:
                            doc.add_picture(tmp_path, width=Inches(5))
                        finally:
                            os.unlink(tmp_path)
                        continue
                p.add_run(f'[image: {img.get("alt", "")}]')
            else:
                _add_inline(p, children)

        elif ttype == "block_code":
            code = tok.get("raw", "").rstrip("\n")
            p = doc.add_paragraph()
            run = p.add_run(code)
            run.font.name = "Courier New"
            run.font.size = Pt(9)

        elif ttype == "list":
            _add_list(doc, tok)

        elif ttype == "block_quote":
            for child in tok.get("children", []):
                if child["type"] == "paragraph":
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    _add_inline(p, child.get("children", []))

        elif ttype == "thematic_break":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("───────────────────────────────")

        elif ttype == "table":
            _add_table(doc, tok)

    return doc


def main():
    args = sys.argv[1:]

    reverse = "-r" in args
    extract_images = "--images" in args
    args = [a for a in args if a not in ("-r", "--images")]

    if not args:
        print("Usage: docx2md [-r] [--images] <input> [output]", file=sys.stderr)
        print("  docx2md input.docx         Convert DOCX → Markdown", file=sys.stderr)
        print("  docx2md -r input.md         Convert Markdown → DOCX", file=sys.stderr)
        sys.exit(1)

    input_file = args[0]
    if not os.path.exists(input_file):
        print(f'Error: File "{input_file}" not found', file=sys.stderr)
        sys.exit(1)

    base, _ = os.path.splitext(input_file)
    default_out_ext = ".docx" if reverse else ".md"
    output_file = args[1] if len(args) > 1 else base + default_out_ext

    if os.path.abspath(input_file) == os.path.abspath(output_file):
        print(
            f'Error: Output file would overwrite input file "{input_file}". '
            "Please specify a different output file.",
            file=sys.stderr,
        )
        sys.exit(1)

    if reverse:
        with open(input_file, "r", encoding="utf-8") as f:
            md_text = f.read()
        doc = convert_md_to_docx(md_text, input_path=input_file)
        doc.save(output_file)
    else:
        image_dir = base + "_images" if extract_images else None
        md = convert(input_file, extract_images=extract_images, image_dir=image_dir)
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(md)
        if extract_images and image_dir and os.path.exists(image_dir):
            count = len(os.listdir(image_dir))
            if count:
                print(f"Extracted {count} image(s) → {image_dir}/")

    print(f"Converted {input_file} → {output_file}")


if __name__ == "__main__":
    main()
